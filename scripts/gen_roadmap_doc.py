# -*- coding: utf-8 -*-
"""
gen_roadmap_doc.py — 生成《Module1 数据清洗标准化 + 归一化字典构造 路线图》Word 文档。

内容严格对齐当前代码：
  biosample_canonical_v2/{layer1_ncbi_attrs,layer2_scan_dump,build_canonical_dict,build_two_part_dict}.py
  Meta2Data/scripts/metadata_downloader.py
  Meta2Data/scripts/column_standardize/column_merge.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(r"Y:\小项目\clone_repo\Meta2Data\Module1_字典与标准化_路线图.docx")

DARK = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
GREY = RGBColor(0x59, 0x59, 0x59)
CODE_BG = "F2F4F7"

doc = Document()

# ── 全局字体 ──
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _set_cn(run, name="微软雅黑"):
    run.font.name = name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name)


def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = DARK
    r.font.size = Pt(17)
    _set_cn(r)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = ACCENT
    r.font.size = Pt(13.5)
    _set_cn(r)
    return p


def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = ACCENT
    r.font.size = Pt(11.5)
    _set_cn(r)
    return p


def body(text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = DARK
        _set_cn(r)
    r = p.add_run(text)
    _set_cn(r)
    return p


def bullet(text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    if lead:
        r = p.add_run(lead)
        r.bold = True
        _set_cn(r)
    r = p.add_run(text)
    _set_cn(r)
    return p


def numbered(text, lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    if lead:
        r = p.add_run(lead)
        r.bold = True
        _set_cn(r)
    r = p.add_run(text)
    _set_cn(r)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p


def rationale(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    r = p.add_run("理由：")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x5A, 0x00)
    _set_cn(r)
    r2 = p.add_run(text)
    r2.font.color.rgb = GREY
    _set_cn(r2)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rr = hdr[i].paragraphs[0].add_run(htext)
        rr.bold = True
        rr.font.size = Pt(9.5)
        _set_cn(rr)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(str(x) for x in row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(val)
            rr.font.size = Pt(9)
            _set_cn(rr)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ══════════════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Meta2Data · Module 1 路线图")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = DARK
_set_cn(tr)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("元数据下载 · 数据清洗与标准化 · BioSample 列名归一化字典构造")
sr.font.size = Pt(12.5)
sr.font.color.rgb = ACCENT
_set_cn(sr)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("绝对细节版 · 每一步算法与理由 · 与代码逐行对齐\n生成日期：2026-07-19")
mr.font.size = Pt(10)
mr.font.color.rgb = GREY
_set_cn(mr)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 0. 本轮改动摘要
# ══════════════════════════════════════════════════════════════════════════
h1("0. 本轮改动与算法升级摘要")

body("升级后的算法（v2）此前已写入 build_canonical_dict.py。本轮工作是把它"
     "在全量数据上跑通并落地：全量重建字典（9.2 万条提取列名）→ 生成两档字典 "
     "→ 部署到 MetaDL 运行时目录 → 用 12 个 BioProject ID 端到端测试通过。"
     "全过程算法自动，未手动清洗任何脏条目（这是明确要求：算法达不到完美也接受，不手改）。")

h3("字典构造算法：v2 相对 v1 的四点升级")
bullet("将候选匹配从「纯 Jaro-Winkler 字符串相似」升级为 "
       "TF-IDF char (2–4)-gram 余弦相似，能捕捉子串/词序无关的形近。", "[A] ")
bullet("融合分 = (TF-IDF 余弦 + Jaro-Winkler) / 2，两个正交信号取平均，"
       "单一信号的偶发高分不足以通过。", "[B] ")
bullet("Jaccard token 过滤：raw 列名与候选官方名必须至少共享一个词，"
       "否则一律否决（消除 collector name → collection_date 这类假阳性）。", "[C] ")
bullet("数据驱动通用词表 + distinctive_name_match：统计每个词在多少个官方名里出现，"
       "高频词（type/name/id/date/host…）视为「通用词」；若两列仅共享通用词、"
       "却各自还带不同的区分性词，则否决（消除 host weight → host_height、"
       "submitter id → submitted_sample_id）。", "[D] ")

h3("运行时标准化：值指纹层（column_merge.py）")
body("除了字典，MetaDL 运行时新增「名字提名 + 值定夺」层：不只看列名，还读列的取值"
     "画像来判两列该不该合。修掉两类错误——名字像但值异义的误合（ReadFile1Md5/"
     "ReadFile2Md5、Platform.1~7 混平台与型号），名字不像但值同义的漏合（Layout↔"
     "LibraryLayout、lat_lon↔LatitudeLongitude）。")

# ══════════════════════════════════════════════════════════════════════════
# 第一部分：字典构造路线
# ══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("第一部分 · BioSample 列名归一化字典构造路线")

body("目标：把 NCBI BioSample 中提交者随手填写的各种原始属性名（attribute_name），"
     "归一到 NCBI 官方标准名（harmonized_name），产出一份「两档」字典供 MetaDL 运行时使用。"
     "整条流水线位于目录 ", )
code("Y:\\小项目\\clone_repo\\biosample_canonical_v2\\")
body("按数据流分五步：Layer 1 官方属性表 → Layer 2 全库 dump 提取 → build_canonical_dict "
     "分型匹配 → build_two_part_dict 分档 → 部署到运行时。")

# ── A1 ──
h2("A1. Layer 1 —— 抓取 NCBI 官方属性表（金标准）")
body("脚本：layer1_ncbi_attrs.py　数据源：", )
code("https://www.ncbi.nlm.nih.gov/biosample/docs/attributes/?format=xml")
h3("算法")
numbered("HTTP 下载官方属性 XML（带 User-Agent，超时 60s）。")
numbered("ElementTree 解析每个 <Attribute>，取 HarmonizedName / Name / "
         "Description / 全部 <Synonym> / <Package>。")
numbered("对每条属性，把 {HarmonizedName, Name} ∪ {synonyms} 合成别名集合 all_aliases。")
numbered("构建 synonym_map：alias（lower-strip）→ harmonized_name，作为精确同义词查表。")
body("产物：", "")
bullet("outputs/layer1_canonical.csv —— harmonized_name | display_name | description | "
       "synonyms(|分隔) | packages（约 960 条官方名）。")
bullet("outputs/layer1_synonym_map.json —— alias → canonical（约 1772 条）。")
rationale("harmonized_name 是 NCBI 机器可读标准键，是整个归一化的锚点（ground truth）。"
          "官方 synonym 表本身就是高质量的人工同义词，直接用作精确命中层，零误判。")

# ── A2 ──
h2("A2. Layer 2 —— 流式扫描全库 dump 提取全部列名")
body("脚本：layer2_scan_dump.py　数据源：", )
code("https://ftp.ncbi.nlm.nih.gov/biosample/biosample_set.xml.gz  （压缩 ~4.4 GB，解压 ~147 GB）")
h3("算法")
numbered("下载 biosample_set.xml.gz 到 downloads/（带进度、断点保留原文件，绝不解压）。")
numbered("gzip.open 流式读，lxml.iterparse 逐 <BioSample> 处理，处理完立即 elem.clear() "
         "释放内存 —— 147 GB 数据常驻内存不超过单条记录。")
numbered("每个 <Attribute> 取 attribute_name（原始列名）与 harmonized_name（可能为空）。")
numbered("以 attribute_name.lower() 为键计频；同时记录每个原始名对应的 harmonized_name 分布。")
numbered("每个原始名取「出现频次最高」的 harmonized_name 作为其官方映射。")
body("产物：", "")
bullet("outputs/layer2_raw_names.csv —— raw_name | harmonized_name | freq"
       "（去重后约 9.24 万条原始列名，按频次降序）。")
bullet("outputs/layer2_longtail.csv —— 其中 harmonized_name 为空的自定义列名（长尾）。")
rationale("提交者填名千奇百怪（collection.date / Collection Date / coll_date …），"
          "全量 dump 是唯一能覆盖真实分布的来源。流式解析是因为解压后 147 GB 无法全载入；"
          "取众数 harmonized_name 是因为同一原始名偶尔被不同人映射到不同官方名，取多数最稳。")

# ── A3 ──
h2("A3. build_canonical_dict.py —— 分型匹配（算法核心）")
body("输入 layer1_canonical.csv + layer2_raw_names.csv，把 9.24 万原始名按置信度分成 "
     "Type 1–4 四型，外加一份低信度参考。关键参数：")
table(
    ["参数", "默认值", "含义"],
    [
        ["FUSION_THRESHOLD", "0.70", "Type 2b 融合分(TF-IDF+JW)/2 下限"],
        ["JACCARD_MIN", "0.0", "token Jaccard 下限（>0 即要求有共享词）"],
        ["JW_PEER", "0.88", "Type 3b JW 匹配 peer 代表阈值"],
        ["NGRAM_SIM", "0.50", "低信度 n-gram 余弦聚类阈值"],
        ["GENERIC_DF_MIN", "15", "词出现在≥15个官方名 = 通用词"],
        ["NAME_KEEP_FUSION", "0.90", "仅共享通用词时的逃生阀融合分"],
        ["TFIDF_BATCH", "5000", "批量向量化大小（控内存）"],
    ],
    widths=[1.8, 1.0, 3.6],
)

h3("预处理（三种键，全程一致）")
bullet("normalize(text)：驼峰拆分 + 非字母数字转空格 + 小写 → 词序列（如 "
       "'CollectionDate' → 'collection date'）。")
bullet("char_fp(text)：小写并删掉所有非字母数字 → 跨格式指纹（'geo: loc: name' 与 "
       "'GeoLocName' 同指纹 'geolocname'）。")
bullet("token_fp(text)：normalize 后的词集合 frozenset（用于 Type 3a 指纹相等判定）。")

h3("Type 1a —— NCBI ground truth（自动，最高置信）")
body("原始名在 dump 里本就带非空 harmonized_name，直接采用。这是官方已经归一好的，无需任何算法。")

h3("Type 1b —— char_fingerprint 精确指纹命中（自动）")
body("原始名的 char_fp 命中官方名的 char_fp 表即归一。")
code("'collection.date' / 'Collection Date' / 'geo: loc: name'  →  collection_date / geo_loc_name")
rationale("record linkage 的标准预处理（Winkler 1990/2006）：只是标点/大小写/驼峰差异的，"
          "本就是同一个名，属确定性命中，零风险自动合并。")

h3("Type 2a —— synonym 官方同义词精确命中（自动）")
body("normalize(raw) 命中 Layer 1 的官方 synonym 表即归一。")
rationale("官方 synonym 是 NCBI 人工维护的等价词表（MetaMap/NCBO/text2term 均以其为首优先级），"
          "精确命中即可信。")

h3("Type 2b —— TF-IDF + JW 融合匹配（推荐，升级核心 [A][B][C][D]）")
body("对既没指纹命中、也没同义词命中的原始名，做「模糊挂官方名」。逐条算法：")
numbered("TF-IDF 余弦：用全部官方名的 normalize 形式拟合 TfidfVectorizer"
         "(analyzer='char_wb', ngram_range=(2,4), sublinear_tf=True)，"
         "对每个 raw 批量算与所有官方名的余弦，取最高分候选。")
numbered("Jaro-Winkler：算 raw 与该候选的 JW 相似（前缀加权，rapidfuzz）。")
numbered("融合：fusion = (TF-IDF余弦 + JW) / 2。")
numbered("Jaccard token 过滤：J(raw词集, 候选词集) 必须 > JACCARD_MIN，即至少共享一个词。")
numbered("区分性门槛 distinctive_name_match：共享了区分性词→接受；一方词集是另一方子集"
         "（缩写/特化，如 temp_c_min ⊆ temp）→接受；仅共享通用词但 fusion≥0.90→接受"
         "（错拼/格式变体 host_heght→host_height）；否则（仅共享通用词且两边各有区分词）→拒绝。")
numbered("同时满足 fusion≥0.70 且 J>0 且 name_ok，才判为 Type 2b 命中。")
code("experimental factor 4 → experimental_factor  (fusion=0.99, tfidf=1.0, jw=0.98, J=0.67)\n"
     "tot_depth_water_col_m → tot_depth_water_col    (fusion=0.99, J=0.80)\n"
     "host weight → host_height                      ✗ 被 [D] 否决（仅共享通用词 host，各带 weight/height）")
rationale("[A] 纯 JW 对「词序/子串」不敏感，字符 n-gram 余弦能补上（text2term/baae119, "
          "Goncalves 2024）。[B] 双信号平均降低单信号偶发高分（arXiv 1903.08206）。"
          "[C] 无共享词的高相似基本是巧合（collector→collection），Jaccard 一刀切掉"
          "（Bilenko 2003, KDD）。[D] 是本轮最实的加固：通用词（name/id/type…）区分度低，"
          "光靠它俩相似会把 host weight 错挂到 host_height，用「数据驱动 df 统计 + 区分词残余」拦住。")

h3("Type 3a —— token 指纹 peer 聚类（组内推荐）")
body("对仍未命中官方名的池子，按 token_fp（词集合）完全相等分组，≥2 成员成一组，"
     "组代表取组内频次最高者。")
code("P0001: {ena-checklist, ena checklist, ena_checklist, ena - checklist}")
rationale("词集合相同、只是分隔/顺序不同的，几乎必为同物（Jaccard=1 的特例，"
          "Rajaraman & Ullman）。官方表里没有它们，只能内部聚，故进「推荐」而非「自动」。")

h3("Type 3b —— JW 匹配 peer 代表（组扩展）")
body("剩余单例用 rapidfuzz 对所有 peer 代表做 JW 匹配，≥JW_PEER(0.88) 则并入该组。")
rationale("把指纹差一点点（拼写小误 ena-last-udate vs ena-last-update）的孤儿挂回已成形的组。")

h3("Type 4 —— standalone 完全没命中（仅格式标准化，不建议合并）")
body("以上全落空的，只做格式规整，不进字典合并逻辑。高频例：'insdc center name'、"
     "'taxonomic identity marker'。")

h3("低信度参考 —— n-gram TF-IDF 聚类（独立输出，不进主字典）")
body("对 Type 4 再做一次 char n-gram TF-IDF 余弦聚类（阈值 0.50，并查集连通），"
     "结果写 low_confidence_ref.csv 供深度人工参考，绝不进主字典。")
rationale("给人留一条「机器都不敢定、但可能相关」的线索，但明确隔离，不污染自动/推荐结果。")

# ── A4 ──
h2("A4. build_two_part_dict.py —— 打包成「两档」字典")
body("把 canonical_dict.json 重组为运行时消费的两档结构，输出 "
     "outputs/biosample_canonical_dict.json：")
table(
    ["档位", "收录内容", "运行时行为"],
    [
        ["direct 直接合并", "Type1（精确/指纹）+ Type2 synonym", "命中即自动改名/合并"],
        ["recommend 推荐", "Type2 fusion≥0.70 + Type3 peer 组", "只列表，用户逐条确认"],
        ["（不进字典）", "Type4 standalone", "仅格式标准化"],
    ],
    widths=[1.5, 2.6, 2.3],
)
bullet("AUTO_FUSION 默认 None：融合分是「模糊相似」而非「清理后相等」，故默认全部进 recommend，"
       "不自动合并（仅显式给 --auto-fusion 阈值时高分才进 direct，作逃生阀）。")
bullet("is_too_short 跳过 normalize 后 ≤3 字符且无空格的极短词（g、otu），避免误命中。")
bullet("recommend 中已被 direct 收录的变体会被剔除（direct 优先）。")
rationale("record linkage 常规分档（Christen 2012《Data Matching》；text2term 亦保留人审界面）："
          "确定性命中自动落地，模糊命中交人审。设计铁律——除 direct 外只出建议、绝不自动合并。")

# ── A5 ──
h2("A5. 部署到 MetaDL 运行时")
body("把 outputs/biosample_canonical_dict.json 复制到 MetaDL 的列标准化目录，"
     "命名为主字典 + 回退字典：")
code(r"Meta2Data\scripts\column_standardize\biosample_canonical_v2.json    ← 主字典" + "\n"
     r"Meta2Data\scripts\column_standardize\biosample_canonical_dict.json  ← 回退（v2 缺失时用）")
body("metadata_downloader.py 第 82 行优先加载 biosample_canonical_v2.json，不存在则回退。")

# ── A6 本次全量结果 ──
h2("A6. 本次全量重建结果（2026-07-19）")
body("对全部 92,413 条原始列名运行，分型统计：")
table(
    ["类型", "条数", "占比 / 说明"],
    [
        ["Type 1 精确命中", "2,051", "1a NCBI=1,920 + 1b 指纹=131"],
        ["Type 2 官方名匹配", "12,611", "2a synonym=76 + 2b 融合=12,535"],
        ["Type 3 peer 组合并", "27,867", "3a 指纹=8,925(3,989组) + 3b JW=18,942"],
        ["Type 4 完全没命中", "49,884", "55.1%，仅格式标准化"],
        ["低信度参考组", "1,911 组", "独立输出，不进主字典"],
    ],
    widths=[2.0, 1.2, 3.2],
)
body("两档字典规模：", "")
bullet("direct：811 canonical 键 / 2,135 变体（Type1=1,248 + synonym=76）。")
bullet("recommend：4,604 canonical 键 / 36,394 变体（fusion=12,516 + peer 3,989 组 23,878 条）。")

# ══════════════════════════════════════════════════════════════════════════
# 第二部分：Module 1 路线
# ══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("第二部分 · Module 1（MetaDL）整体路线")

body("入口：bin/Meta2Data-MetaDL（bash 包装）→ scripts/metadata_downloader.py。"
     "输入一个装 BioProject ID 的目录（每行一个，支持 PRJNA/PRJEB/PRJDB/PRJCA 等），"
     "输出一张跨库合并、清洗、标准化后的元数据总表及配套产物。")

h2("B1. 输入读取与 ID 分类")
numbered("read_input_ids：扫描输入目录下所有 .txt，逐行读 ID，去重。")
numbered("classify_ids / ID_PATTERNS：正则分流——PRJC*→CNCB，PRJ[EDN]*→NCBI，"
         "SAM*→BioSample，[CEDS]R[RSXP]*→SRA accession。")

h2("B2. 下载（每个 BioProject 独立处理，可并行 + 断点续跑）")
body("process_single_bioproject 按来源分派：")
h3("NCBI 分支（download_ncbi_metadata）")
numbered("由 BioProject 查其全部 BioSample，下载 BioSample 属性（parse_biosample_file）。")
numbered("查 SRA，下载 SRA RunInfo（Run 级粒度）。")
numbered("merge_ncbi_data_single：以 SRA RunInfo 为左表（Run 级），"
         "BioSample 属性按 Biosample 键 left join（回退 Run 键）。")
h3("CNCB/GSA 分支（download_cncb_metadata）")
numbered("经 NGDC HTTP API 取 GSA 元数据表（tsv/xlsx）。")
body("公共后处理：", "")
bullet("保留用户原始 BioProject ID（NCBI 可能把 PRJEB/PRJDB 交叉引用成 PRJNA）。")
bullet("附加 Description（预取缓存，避免并行触发 NCBI 429）与 Source_Database。")
bullet("validate_run_data：无有效 Run 信息的项目标记失败、不进合并。")
bullet("每项目落盘 <ID>.processed.csv（断点：已完成的重跑直接读缓存）。")
rationale("以 Run 为最细粒度左表，保证「一条测序 Run 一行」；BioSample 属性可一对多广播到"
          "同 BioSample 的多个 Run。逐项目落盘 + StateManager 使 4.4GB 级任务可中断续跑。")

h2("B3. 单项目级清洗与标准化（standardize_columns）")
body("在每个项目 DataFrame 上先做一轮轻量清洗（注意：字典同义词合并已【移出】此步，"
     "改为在全表合并后统一做，好让共现规则看到整个数据集）。")
h3("① clean_and_standardize_columns —— 列名清洗 + 内容探测")
numbered("remove_empty_columns：删全空列（含只有 nan/空白的列）。")
numbered("remove_duplicate_columns：删内容逐格完全相同的重复列。")
numbered("去前后缀：^(Sample|Experiment|Run)_ 前缀、(_sra|_biosample|_x|_y)$ 后缀。")
numbered("首列若叫 'ID' 直接丢弃（无意义索引）。")
numbered("内容探测 UNIFIED_PATTERNS：若某列所有非空值都匹配某模式，则改名为标准列——"
         "Run(^[CEDS]RR\\d+$)、Bioproject(^PRJ[CEDN][A-Z]\\d+$)、"
         "Biosample(^SAM[CEDN][A-Z]?\\d+$)、Experiment(^[CEDS]RX\\d+$)。")
numbered("_apply_priority_order：PRIORITY_COLUMNS（Run/Bioproject/Description/"
         "DesignDescription/Biosample/Experiment）提到最前。")
rationale("内容探测（看值而非看名）能把任意乱名的 ID 列认出来，比纯靠列名鲁棒得多；"
          "标准 ID 是下游跨库对齐的主键，必须先立起来。")
h3("② apply_camelcase_normalization —— 驼峰归一 + 智能并列")
numbered("含空格/下划线/连字符的列名转 CamelCase（library_strategy → LibraryStrategy）。")
numbered("若归一后撞上已存在的同名列，则按行智能合并单元格（都空→空；一空→取非空；"
         "相等→取其一；都非空且不同→'a_b'），然后删源列。")
h3("③ DesignDescription 别名归一")
body("把 CNCB 侧的 LibraryConstruction/ExperimentalDesign 等 6 种别名统一改名为 DesignDescription。")

h2("B4. 全表最终合并（merge_all_results，核心清洗管线）")
body("把所有 <ID>.processed.csv 纵向 concat 后，按 a–k 顺序清洗、标准化、产表：")
numbered("a. 分离无 Run 记录 → RecordWithoutRUNinfo.csv，主表剔除。", )
numbered("b. 删下载/路径列：DownloadPath、DownloadReadFile1/2、ReadFilename1/2。")
numbered("c. CNCB 去重（deduplicate_cncb_columns）：清掉 CNCB 列中与同行核心列取值相同的冗余单元格。")
numbered("d. 删全空列。")
numbered("e. LatLon 解析：把 '32.4 N 119.4 E' 拆成数值 lat/lon（度符号/逗号清理，"
         "缺失填补而非覆盖），删原 LatLon。")
numbered("f. FileSize→SizeMb：字节转兆（支持 'a|b' 求和），仅填补 SizeMb 空缺，删原 FileSize。")
numbered("g. 【字典 direct 档】column_merge.apply_direct：在全表上一次性应用直接合并档"
         "（char_fp 命中即改名/合并到 canonical），protected=CORE_COLUMNS∪{Source_Database} 永不改名。")
numbered("h. 列排序：核心列（CORE_COLUMNS）在前，其余按 CamelCase 前缀分组、组内外字母序。")
numbered("i. generate_column_description：每列出 类型/非空数/填充率/覆盖数据集数/Top5 值。")
numbered("j. 写 all_metadata_merged.csv（utf-8-sig）。")
numbered("k. 【字典 recommend 档 + 值指纹层】column_merge.build_recommend_table：只产建议、"
         "不改数据，写 merge_recommendations.csv + merge_groups.txt。")
rationale("字典合并从「逐项目」上移到「全表一次」，是为了让 recommend 的『共现≥2 成员才提示』"
          "规则能看到整个数据集的列共现，判断更准；direct 用 char_fp 等价类实现跨数据集统一命名。")

h2("B5. 运行时标准化两层（column_merge.py 详解）")
h3("direct 直接合并档（自动落地）")
body("列的 char_fp 命中 direct 档即改名/合并到 canonical，把 'GeoLocName'、'geo loc name' "
     "统一到 'geo_loc_name'。protected 核心列绝不改名（防 HostTaxonomyId→host_taxid 误伤下游）。")
h3("recommend 推荐档：三个候选来源（只出表，不改数据）")
bullet("dict_recommend（字典第一层）：命中 recommend 档、且目标 canonical 在本表共现成员≥2 的列。", )
bullet("internal_fresh（数据内部）：未命中字典的【文本类】列，用 TF-IDF+JW 融合(≥0.70)"
       "挂最近官方名；跳过 protected/数字/日期/坐标/ID 类（修 StudyPubmedId→study_name 误报）。", )
bullet("value_peer（值指纹层，名字提名+值定夺）：见下。", )
h3("值指纹层算法（value_aware_peer_groups + column_relation）")
numbered("value_profile：给每列算值画像 {kind, vset(≤400 值), card, avg_tokens}，"
         "kind 由 _cell_kind 粗判：md5 / hexhash / accession / geo_coord（须带小数）/ "
         "date / numeric / text / free_text(avg_tokens≥4)。")
numbered("提名：列对名字相似度≥NAME_PROPOSE_THR(0.55) 或共享词 token 才成候选；"
         "且每条边至少一个非 protected 列（不合并两个核心 canonical）。")
numbered("定夺 column_relation：①共现（共同有值行占比≥COOCCUR_MIN_FRAC=0.15）→ 看行内一致率，"
         "≥COOCCUR_AGREE_MERGE(0.60) 判 merge，否则 veto（平行字段）；②填充互补 → "
         "geo_coord 直接 merge；text/free_text 需词表 Jaccard≥TEXT_JACCARD_MERGE(0.34)；"
         "md5/hexhash/accession/numeric/date 值帮不上忙 → weak。")
numbered("并查集连通成组，代表名优先取 protected/标准列，否则取填充率最高者。")
code("Layout ↔ LibraryLayout           ✓ merge（值同义，名字不像也挖出）\n"
     "lat_lon ↔ LatitudeLongitude       ✓ merge（geo_coord 互补）\n"
     "Platform.1/.3/.5/.7 vs .2/.4/.6    ✓ 正确拆两组（值冲突→veto，不混平台与型号）\n"
     "ReadFile1Md5 vs ReadFile2Md5       ✗ 不合（md5 值证据→weak，名字严判）")
rationale("只比列名会同时误合（名字像值异义）和漏合（名字不像值同义）。让『名字负责提名、"
          "值负责定夺』：共现看一致率区分重复列 vs 平行字段；互补填充只在值能证明字段身份"
          "（坐标格式特异、文本词表大面积重叠）时才敢合；ID/数值类值集合无判别力，退回名字严判。"
          "阈值都是踩过坑校准的（如坐标正则强制带小数，否则 Title 里 '16s of B10' 被误判经纬度）。")

h3("Phase B：用户确认后落地（--apply-merges）")
body("用户编辑 merge_groups.txt（每行『目标/成员1/成员2』），"
     "run_apply_merges → apply_user_groups 按分组把列合并落地，"
     "写 all_metadata_merged.standardized.csv。整个合并只有 direct 档自动，其余都需人点头。")

h2("B6. 产物清单")
table(
    ["产物", "内容"],
    [
        ["all_metadata_merged.csv", "最终合并总表（已应用 direct 档字典合并）"],
        ["all_metadata_merged.standardized.csv", "--apply-merges 落地用户确认合并后的表"],
        ["merge_recommendations.csv", "推荐合并候选表（逗号分隔，每列独立；dict/internal_fresh/value_peer 三源）"],
        ["merge_groups.txt", "可编辑分组模板（删你不要的行，再 --apply-merges）"],
        ["status.tsv", "每个输入 ID 的下载状态"],
        ["column_description.tsv", "每列 类型/非空/填充率/覆盖数据集/Top5 值"],
        ["bioproject_absdesc.tsv", "每个 BioProject 的 PMID/PMC/DOI/摘要"],
        ["RecordWithoutRUNinfo.csv", "无 SRA Run 信息的记录（有则输出）"],
        ["tmp/ 目录", "断点续跑状态 + 各 BioProject 中间产物（可删）"],
    ],
    widths=[2.6, 3.8],
)

h2("B7. 本轮端到端测试结果")
body("用两组共 12 个 BioProject（跨 NCBI/CNCB/DDBJ/ENA）直接跑 metadata_downloader.py：")
table(
    ["测试集", "结果", "direct 合并", "记录/列", "推荐候选"],
    [
        ["input2（4 ID）", "4/4 成功", "14 列", "68 记录 / 65 列",
         "project_name←ProjectID；heavy_metals_meth←HeavyMetalsMethod"],
        ["input1（8 ID）", "8/8 成功", "10 列", "360 记录 / 74 列",
         "description←Design/PublicDescription；Strategy/Layout/lat_lon 自动挖出；Platform 拆两组；Md5 无误合"],
    ],
    widths=[1.3, 0.9, 0.9, 1.3, 2.6],
)
body("注（与字典无关的环境项）：bin 包装脚本用 python3 调用，在本机解析异常会报 exit 49，"
     "故测试直接用完整 Python 路径调 metadata_downloader.py；CRA*.temp.xlsx 偶发 WinError 32 "
     "文件占用为 Windows 临时文件竞争，8/8 仍全部成功。")

doc.save(OUT)
print(f"[written] {OUT}")
print(f"  size = {OUT.stat().st_size:,} bytes")
